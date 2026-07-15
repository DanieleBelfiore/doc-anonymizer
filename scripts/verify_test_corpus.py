"""Verify test_output/ against expectations for each case in test_input/.

Run engine/processor.py first, then this script:
    engine/venv/bin/python -m engine.processor --input test_input --output test_output
    engine/venv/bin/python scripts/verify_test_corpus.py

For PDFs we compare extracted text; for docx, paragraph+cell text; for txt,
raw content. Each entry in CHECKS lists strings that must be gone from the
output (real PII) and strings that must remain (false-positive bait / non-PII
context). Exits non-zero if any case fails, so it can be wired into CI later.

Requires an Ollama sidecar with the configured model running — detection
uses live LLM calls, so results can drift if the model or prompt changes.
"""

import sys
from pathlib import Path

import fitz
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "test_output"


def pdf_text(path: Path) -> str:
    doc = fitz.open(path)
    text = "\n".join(page.get_text("text") for page in doc)
    doc.close()
    return text


def docx_text(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# filename -> (must_be_gone: [strings], must_remain: [strings])
CHECKS: dict[str, tuple[list[str], list[str]]] = {
    "01_clinical_standard_dob_before.pdf": (
        ["Mario Ferrari", "14/03/1978", "FRRMRA78C14F205X", "Via Roma 22"],
        ["Glicemia", "92 mg/dL"],
    ),
    "02_clinical_dob_after_label.pdf": (
        ["COLOMBO", "LAURA", "09/11/1980"],
        ["Colesterolo Totale"],
    ),
    "03_clinical_multidate_competing.pdf": (
        ["ESPOSITO", "MARCO", "18/09/1973"],
        ["Trigliceridi"],
    ),
    "04_allcaps_name_column_spacing.pdf": (
        ["RIZZO", "GIUSEPPINA"],
        ["Cardiologia", "ECG"],
    ),
    "05_false_positive_bait.pdf": (
        [],
        ["HGB", "RBC", "GLU", "MCV", "MCH", "MCHC", "HCT", "IVA 22%", "Cardiologia", "3001", "PR-2026-00891"],
    ),
    "06_commercial_letter_piva.pdf": (
        ["Alberto Conti", "04521870582"],
        ["2.350,00", "2026/118"],
    ),
    "07_iban_document.pdf": (
        ["Sara Moretti", "IT28W0300203280123456789012"],
        ["118/2026", "480,00"],
    ),
    "08_id_document_number.pdf": (
        ["Elena Fontana", "CA9284712"],
        ["Comune di Torino"],
    ),
    "09_license_plate_near_wrong_date.pdf": (
        ["Davide Serra", "GH482LM"],
        [],
    ),
    "10_license_plate_dob_stress.pdf": (
        ["Paolo Gatti", "BC901XY"],
        [],
    ),
    "11_address_person_name_collision.pdf": (
        ["Carlo Bianchi", "via Carlo Bianchi 5", "Roma Verdi", "piazza Roma 7"],
        [],
    ),
    "12_isolated_city_no_redact.pdf": (
        [],
        ["Roma e' la capitale", "Bologna e Milano"],
    ),
    "13_multi_person_document.pdf": (
        ["Federica Longo", "Stefano Ricci", "Chiara Bruno", "Luca Ferri"],
        [],
    ),
    "14_email_phone_simple.pdf": (
        ["assistenza@example-clinic.it", "349-2214578"],
        [],
    ),
    "15_no_pii_document.pdf": (
        [],
        ["XN-1000", "calibrazione", "controllo qualita'"],
    ),
    "16_single_line_short.pdf": (
        [],
        ["Nessuna informazione sensibile"],
    ),
    "17_multipage_long.pdf": (
        ["Tommaso Villa", "VLLTMS75M20L219K", "Via dei Mille 40", "Enrico Pace", "010-5544332"],
        [],
    ),
    "18_scanned_ocr_fallback.pdf": (
        ["Ilaria Bruno", "BRNLRI88T45G273P"],
        [],
    ),
    "19_accented_characters.pdf": (
        ["Perla D'Angelo"],
        ["CITTÀ DI PERUGIA", "Città della Pieve", "Perché il referto"],
    ),
    "20_repeated_name_occurrences.pdf": (
        ["Simone Galli", "Anna Galli"],
        [],
    ),
    "docx_01_name_split_across_runs.docx": (
        ["Vittorio", "Testa"],
        ["14/02/2026"],
    ),
    "docx_02_table_redaction.docx": (
        ["Silvia Conte", "CNTSLV82D55H501R", "Massimo Riva", "RIVMSM79A01F205Y"],
        [],
    ),
    "docx_03_empty_paragraphs.docx": (
        ["Nadia Costa"],
        ["Nessuna anomalia"],
    ),
    "docx_04_repeated_name_multi_cell.docx": (
        ["Beatrice Marchetti"],
        ["20/01/2026"],
    ),
    "docx_05_name_across_paragraphs.docx": (
        ["FILIPPETTI", "GIANMARCO"],
        ["Ortopedia"],
    ),
    "21_multiline_entity_name.pdf": (
        ["MARTINELLI", "ALESSANDRO", "10/10/1990"],
        ["Ortopedia"],
    ),
    "22_iban_spaced.pdf": (
        ["Roberto Fabbri", "IT60 X054 2811 1010 0000 0123 456"],
        [],
    ),
    "24_foreign_name.pdf": (
        ["John Smith", "Marie Dubois"],
        [],
    ),
    "25_pec_intl_phone.pdf": (
        ["mario.bianchi@pec.example.it", "+39 340 1122334"],
        ["2026/774"],
    ),
    "26_lowercase_cf.pdf": (
        ["Luigi Amato", "mtalgu85r10f839w"],
        [],
    ),
    "test.txt": (
        ["Mario Rossi", "mario@gmail.com"],
        [],
    ),
    "test2.txt": (
        ["Giuseppe Verdi", "IT60X0542811101000000123456", "AV1234567", "FX123GH", "12.05.1985"],
        [],
    ),
}

# Files the engine is expected to SKIP (per-file warning, no output written).
EXPECT_NO_OUTPUT = {
    "23_encrypted_pdf.pdf",  # password-protected: unreadable without the password
}


def main() -> int:
    results = []

    for filename in sorted(EXPECT_NO_OUTPUT):
        out_path = OUT / filename
        if out_path.exists():
            results.append((filename, "FAIL", ["expected engine to skip this file, but output exists"]))
        else:
            results.append((filename, "PASS", []))

    for filename, (must_be_gone, must_remain) in CHECKS.items():
        out_path = OUT / filename
        if not out_path.exists():
            results.append((filename, "MISSING OUTPUT", []))
            continue

        if filename.endswith(".pdf"):
            text = pdf_text(out_path)
        elif filename.endswith(".docx"):
            text = docx_text(out_path)
        else:
            text = out_path.read_text(encoding="utf-8")

        problems = []
        for s in must_be_gone:
            if s in text:
                problems.append(f"LEAK: {s!r} still present")
        for s in must_remain:
            if s not in text:
                problems.append(f"FALSE POSITIVE: {s!r} was removed but should remain")

        results.append((filename, "PASS" if not problems else "FAIL", problems))

    fails = [r for r in results if r[1] != "PASS"]
    print(f"\n{len(results) - len(fails)}/{len(results)} cases passed\n")
    for filename, status, problems in results:
        marker = "OK  " if status == "PASS" else "FAIL"
        print(f"[{marker}] {filename}")
        for p in problems:
            print(f"       {p}")

    return 1 if fails else 0


if __name__ == "__main__":
    sys.exit(main())
