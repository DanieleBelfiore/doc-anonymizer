"""Generate a synthetic test corpus for engine/processor.py.

All PII in these documents is fabricated. Run from the repo root:
    engine/venv/bin/python scripts/generate_test_corpus.py

Regenerates test_input/ from scratch — each case exercises a distinct
detection or format-redaction mechanism (see the case list below), rather
than being random variations of the same scenario.
"""

import fitz  # PyMuPDF
from docx import Document
from PIL import Image, ImageDraw, ImageFont
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "test_input"
FONT_PATH = "/System/Library/Fonts/Supplemental/Arial.ttf"


def write_text_pdf(filename: str, lines: list[str], fontsize: float = 11) -> None:
    doc = fitz.open()
    page = doc.new_page()
    y = 50
    for line in lines:
        page.insert_text((50, y), line, fontsize=fontsize, fontname="helv")
        y += fontsize + 6
        if y > 780:
            page = doc.new_page()
            y = 50
    doc.save(OUT / filename)
    doc.close()


def write_multipage_pdf(filename: str, pages: list[list[str]], fontsize: float = 11) -> None:
    doc = fitz.open()
    for lines in pages:
        page = doc.new_page()
        y = 50
        for line in lines:
            page.insert_text((50, y), line, fontsize=fontsize, fontname="helv")
            y += fontsize + 6
    doc.save(OUT / filename)
    doc.close()


def write_scanned_pdf(filename: str, lines: list[str]) -> None:
    """Render text as a flat image with no text layer, forcing the OCR fallback path."""
    img = Image.new("RGB", (1240, 1754), "white")  # ~A4 at 150dpi
    draw = ImageDraw.Draw(img)
    font = ImageFont.truetype(FONT_PATH, 28)
    y = 80
    for line in lines:
        draw.text((80, y), line, fill="black", font=font)
        y += 45

    doc = fitz.open()
    page = doc.new_page(width=1240, height=1754)
    page.insert_image(page.rect, stream=_img_to_bytes(img))
    doc.save(OUT / filename)
    doc.close()


def _img_to_bytes(img: Image.Image) -> bytes:
    import io
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def write_encrypted_pdf(filename: str, lines: list[str]) -> None:
    """Password-protected PDF: the engine can't read it and must emit a
    per-file warning instead of crashing the whole batch."""
    doc = fitz.open()
    page = doc.new_page()
    y = 50
    for line in lines:
        page.insert_text((50, y), line, fontsize=11, fontname="helv")
        y += 17
    doc.save(
        OUT / filename,
        encryption=fitz.PDF_ENCRYPT_AES_256,
        user_pw="segreta",
        owner_pw="segreta",
    )
    doc.close()


def write_docx(filename: str, builder) -> None:
    doc = Document()
    builder(doc)
    doc.save(OUT / filename)


# ---- PDF cases -------------------------------------------------------

PDF_CASES: dict[str, list[str]] = {
    "01_clinical_standard_dob_before.pdf": [
        "CENTRO MEDICO SAN LORENZO",
        "Referto di Laboratorio",
        "",
        "Paziente: Mario Ferrari",
        "Data di nascita: 14/03/1978",
        "CF: FRRMRA78C14F205X",
        "Indirizzo: Via Roma 22, 40121 Bologna",
        "",
        "Esame: Glicemia - 92 mg/dL - Intervallo: 70-100",
    ],
    "02_clinical_dob_after_label.pdf": [
        "OSPEDALE SANTA CHIARA",
        "",
        "Sig.  COLOMBO   LAURA",
        "Id.: 5512309",
        "Eta': 45 Anni",
        "09/11/1980",
        "Data Nascita:",
        "",
        "Esame: Colesterolo Totale - 195 mg/dL",
    ],
    "03_clinical_multidate_competing.pdf": [
        "LABORATORIO ANALISI CENTRALE",
        "Data di Refertazione",
        "02/02/2026",
        "",
        "Sig.  ESPOSITO   MARCO",
        "Eta':  52  Anni",
        "18/09/1973",
        "Data Nascita:",
        "",
        "Richiedente: Dr. Sig.  ESPOSITO   MARCO",
        "Richiesta:",
        "30/01/2026",
        "Prelievo ore: 08:15",
        "del 30/01/2026",
        "Check-In:",
        "",
        "Esame: Trigliceridi - 140 mg/dL",
        "Firmatario: Dr. Anna Neri  Data e ora della firma: 02/02/2026 10:12:00.",
    ],
    "04_allcaps_name_column_spacing.pdf": [
        "AZIENDA OSPEDALIERA UNIVERSITARIA",
        "",
        "RIZZO      GIUSEPPINA",
        "Sig.ra",
        "Reparto: Cardiologia",
        "",
        "Richiedente: Sig.ra  RIZZO      GIUSEPPINA",
        "",
        "Esame: ECG - Nella norma",
    ],
    "05_false_positive_bait.pdf": [
        "REFERTO ESAMI EMATOCHIMICI",
        "",
        "Esami: HGB 14.2 g/dL, RBC 5.1 x10^6/uL, GLU 90 mg/dL,",
        "MCV 82.9 fL, MCH 28.4 pg, MCHC 34.3 g/dL, HCT 47.5%",
        "IVA 22% applicata sui referti privati.",
        "Reparto: Cardiologia",
        "Tipo Reparto: ESTERNI",
        "Accettazione Num. 3001",
        "Codice pratica: PR-2026-00891",
        "Nota: valori nella norma per soggetti adulti.",
    ],
    "06_commercial_letter_piva.pdf": [
        "Studio Legale Bianchi & Associati",
        "Via Torino 8, 00185 Roma",
        "",
        "Spett.le Cliente,",
        "in riferimento alla pratica n. 2026/118 intestata al",
        "Sig. Alberto Conti (P.IVA 04521870582), si comunica",
        "che il pagamento di euro 2.350,00 e' stato registrato.",
        "",
        "Cordiali saluti,",
        "Studio Legale Bianchi & Associati",
    ],
    "07_iban_document.pdf": [
        "DISPOSIZIONE DI BONIFICO",
        "",
        "Ordinante: Sara Moretti",
        "IBAN: IT28W0300203280123456789012",
        "Causale: Saldo fattura 118/2026",
        "Importo: euro 480,00",
    ],
    "08_id_document_number.pdf": [
        "MODULO DI IDENTIFICAZIONE OSPITE",
        "",
        "Nome e Cognome: Elena Fontana",
        "Documento: Carta d'Identita' n. CA9284712",
        "Rilasciato da: Comune di Torino",
        "",
        "Motivo visita: Consulenza ambulatoriale",
    ],
    "09_license_plate_near_wrong_date.pdf": [
        "VERBALE DI DEPOSITO VEICOLO",
        "",
        "Proprietario: Davide Serra",
        "Targa: GH482LM",
        "Data deposito: 05/04/2026",
        "nato il 21/11/1990",
        "",
        "Ritiro consentito previa esibizione documento.",
    ],
    "10_license_plate_dob_stress.pdf": [
        "CONTRATTO DI NOLEGGIO",
        "",
        "Conducente: Paolo Gatti, nato il 03/03/1985",
        "Targa veicolo: BC901XY",
        "Data ritiro: 12/12/2026  Data riconsegna: 19/12/2026",
    ],
    "11_address_person_name_collision.pdf": [
        "SCHEDA ANAGRAFICA",
        "",
        "Carlo Bianchi, paziente, visitato presso ambulatorio",
        "sito in via Carlo Bianchi 5, Torino.",
        "",
        "La sig.ra Roma Verdi vive in piazza Roma 7, Bologna.",
    ],
    "12_isolated_city_no_redact.pdf": [
        "NOTA INFORMATIVA",
        "",
        "Roma e' la capitale d'Italia.",
        "Il congresso si terra' a Bologna e Milano.",
        "Nessun dato personale in questo documento.",
    ],
    "13_multi_person_document.pdf": [
        "REFERTO CONSULTO MULTIDISCIPLINARE",
        "",
        "Paziente: Federica Longo",
        "Medico curante: Dr. Stefano Ricci",
        "Medico consulente: Dr.ssa Chiara Bruno",
        "Infermiere responsabile: Luca Ferri",
    ],
    "14_email_phone_simple.pdf": [
        "CONTATTI UFFICIO",
        "",
        "Per informazioni scrivere a: assistenza@example-clinic.it",
        "oppure telefonare al numero 349-2214578.",
    ],
    "15_no_pii_document.pdf": [
        "MANUALE OPERATIVO INTERNO",
        "",
        "Procedura di calibrazione strumento XN-1000.",
        "Verificare che il valore di controllo qualita'",
        "rientri nell'intervallo di riferimento previsto.",
        "Nessun dato paziente e' presente in questo documento.",
    ],
    "16_single_line_short.pdf": [
        "Nessuna informazione sensibile qui dentro.",
    ],
    "19_accented_characters.pdf": [
        "OSPEDALE CITTÀ DI PERUGIA",
        "",
        "Paziente: Perla D'Angelo",
        "Nato/a a: Città della Pieve, è residente",
        "presso l'abitazione di Via dell'Ospedale 3.",
        "Perché il referto è urgente, sarà inviato più tardi.",
    ],
    "20_repeated_name_occurrences.pdf": [
        "CARTELLA CLINICA",
        "",
        "Paziente: Simone Galli",
        "Il paziente Simone Galli e' stato ricoverato in data 01/02/2026.",
        "Referente familiare del paziente Simone Galli: Anna Galli.",
        "Dimissione firmata a nome di Simone Galli.",
    ],
    # Column-layout extraction can split a single name across two lines;
    # the LLM reproduces it with a plain space, so offset resolution must
    # match across the newline.
    "21_multiline_entity_name.pdf": [
        "SCHEDA ACCETTAZIONE",
        "",
        "Paziente:",
        "MARTINELLI",
        "ALESSANDRO",
        "Data di nascita: 10/10/1990",
        "Reparto: Ortopedia",
    ],
    # Italian IBANs are commonly printed in space-separated groups of 4;
    # the LLM may return the compact form, so matching must bridge the gaps.
    "22_iban_spaced.pdf": [
        "COORDINATE BANCARIE",
        "",
        "Intestatario: Roberto Fabbri",
        "IBAN: IT60 X054 2811 1010 0000 0123 456",
        "Banca: Istituto di Credito Esempio",
    ],
    "24_foreign_name.pdf": [
        "REFERTO AMBULATORIALE",
        "",
        "Il paziente John Smith, nato a Londra, e' stato",
        "visitato presso il nostro ambulatorio.",
        "Accompagnatore: Marie Dubois.",
    ],
    "25_pec_intl_phone.pdf": [
        "CONTATTI PRATICA",
        "",
        "PEC: mario.bianchi@pec.example.it",
        "Telefono: +39 340 1122334",
        "Riferimento pratica: 2026/774",
    ],
    "26_lowercase_cf.pdf": [
        "MODULO PRIVACY",
        "",
        "Firmatario: Luigi Amato",
        "cf: mtalgu85r10f839w",
        "Il presente modulo autorizza il trattamento dei dati.",
    ],
}

MULTIPAGE_CASE = (
    "17_multipage_long.pdf",
    [
        [
            "REFERTO MULTI-PAGINA - Pagina 1",
            "",
            "Paziente: Tommaso Villa",
            "CF: VLLTMS75M20L219K",
        ],
        [
            "Pagina 2 - Anamnesi",
            "",
            "Nessun dato aggiuntivo rilevante in questa sezione.",
            "Indirizzo: Via dei Mille 40, 16122 Genova",
        ],
        [
            "Pagina 3 - Conclusioni",
            "",
            "Contatto medico curante: Dr. Enrico Pace",
            "Tel: 010-5544332",
        ],
    ],
)

SCANNED_CASE = (
    "18_scanned_ocr_fallback.pdf",
    [
        "REFERTO SCANSIONATO",
        "Paziente: Ilaria Bruno",
        "CF: BRNLRI88T45G273P",
        "Data di nascita: 05/12/1988",
    ],
)


# ---- DOCX cases -------------------------------------------------------

def docx_name_split_across_runs(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run("Il paziente ")
    p.add_run("Vittorio").bold = True
    p.add_run(" ")
    p.add_run("Testa").bold = True
    p.add_run(" e' stato visitato in data 14/02/2026.")


def docx_table_redaction(doc: Document) -> None:
    doc.add_paragraph("Elenco pazienti reparto Cardiologia")
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Nome"
    table.rows[0].cells[1].text = "CF"
    table.rows[1].cells[0].text = "Silvia Conte"
    table.rows[1].cells[1].text = "CNTSLV82D55H501R"
    table.rows[2].cells[0].text = "Massimo Riva"
    table.rows[2].cells[1].text = "RIVMSM79A01F205Y"


def docx_empty_paragraphs(doc: Document) -> None:
    doc.add_paragraph("Referto ambulatoriale")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("Paziente: Nadia Costa")
    doc.add_paragraph("")
    doc.add_paragraph("Nessuna anomalia riscontrata.")


def docx_repeated_name_multi_cell(doc: Document) -> None:
    doc.add_paragraph("Registro accessi - Beatrice Marchetti")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Data"
    table.rows[0].cells[1].text = "Paziente"
    table.rows[1].cells[0].text = "20/01/2026"
    table.rows[1].cells[1].text = "Beatrice Marchetti"


def docx_name_across_paragraphs(doc: Document) -> None:
    # Name split across two consecutive paragraphs (column-layout import):
    # the LLM sees "FILIPPETTI\nGIANMARCO" in the joined text and returns the
    # combined name — the entity spans a paragraph boundary and must be
    # clamped per-paragraph, not dropped.
    doc.add_paragraph("SCHEDA ACCETTAZIONE")
    doc.add_paragraph("Paziente:")
    doc.add_paragraph("FILIPPETTI")
    doc.add_paragraph("GIANMARCO")
    doc.add_paragraph("Reparto: Ortopedia")


DOCX_CASES = {
    "docx_01_name_split_across_runs.docx": docx_name_split_across_runs,
    "docx_02_table_redaction.docx": docx_table_redaction,
    "docx_03_empty_paragraphs.docx": docx_empty_paragraphs,
    "docx_04_repeated_name_multi_cell.docx": docx_repeated_name_multi_cell,
    "docx_05_name_across_paragraphs.docx": docx_name_across_paragraphs,
}


# ---- TXT cases ----------------------------------------------------------

TXT_CASES = {
    "test.txt": "Ciao, sono Mario Rossi. Mail: mario@gmail.com",
    "test2.txt": (
        "Giuseppe Verdi, IBAN IT60X0542811101000000123456, \n"
        "documento CI AV1234567, targa FX123GH, nato il 12.05.1985."
    ),
}


def main() -> None:
    for f in OUT.glob("*"):
        if f.name != ".DS_Store" and f.name != ".gitkeep":
            f.unlink()

    for filename, lines in PDF_CASES.items():
        write_text_pdf(filename, lines)

    write_multipage_pdf(*MULTIPAGE_CASE)
    write_scanned_pdf(*SCANNED_CASE)
    write_encrypted_pdf("23_encrypted_pdf.pdf", [
        "DOCUMENTO RISERVATO",
        "Paziente: Nome Protetto",
    ])

    for filename, builder in DOCX_CASES.items():
        write_docx(filename, builder)

    for filename, content in TXT_CASES.items():
        (OUT / filename).write_text(content, encoding="utf-8")

    total = sum(1 for f in OUT.glob("*") if f.name not in (".DS_Store", ".gitkeep"))
    print(f"Generated {total} files in {OUT}")


if __name__ == "__main__":
    main()
