import os
import sys
import argparse
import json
from pathlib import Path
from docx import Document
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
from engine.anonymizer import DocumentAnonymizer

# When running as a PyInstaller bundle, point pytesseract at the bundled binary
if getattr(sys, 'frozen', False):
    _base = sys._MEIPASS  # type: ignore[attr-defined]
    _tess_bin = 'tesseract.exe' if sys.platform == 'win32' else 'tesseract'
    pytesseract.pytesseract.tesseract_cmd = os.path.join(_base, _tess_bin)
    os.environ['TESSDATA_PREFIX'] = os.path.join(_base, 'tessdata')

class FolderProcessor:
    def __init__(self, anonymizer: DocumentAnonymizer, mode="default"):
        self.anonymizer = anonymizer
        self.mode = mode
        self.supported_extensions = {".txt", ".docx", ".pdf"}

    def process_folder(self, input_path: str, output_path: str):
        input_dir = Path(input_path)
        output_dir = Path(output_path)

        if not input_dir.exists() or not input_dir.is_dir():
            print(json.dumps({"status": "error", "message": f"Input folder not found: {input_path}"}))
            sys.stdout.flush()
            return

        output_dir.mkdir(parents=True, exist_ok=True)

        files = [f for f in input_dir.iterdir() if f.suffix.lower() in self.supported_extensions]
        total_files = len(files)
        
        if total_files == 0:
            print(json.dumps({"status": "error", "message": "No supported files found"}))
            return

        for i, file_path in enumerate(files):
            try:
                target_path = output_dir / file_path.name
                if file_path.suffix.lower() == ".txt":
                    self._process_txt(file_path, target_path)
                elif file_path.suffix.lower() == ".docx":
                    self._process_docx(file_path, target_path)
                elif file_path.suffix.lower() == ".pdf":
                    self._process_pdf(file_path, target_path)
                
                progress = int(((i + 1) / total_files) * 100)
                print(json.dumps({
                    "status": "progress", "current": i + 1, "total": total_files, 
                    "percentage": progress, "file": str(file_path.name)
                }))
                sys.stdout.flush()
            except Exception as e:
                print(json.dumps({"status": "warning", "file": str(file_path.name), "error": str(e)}))
                sys.stdout.flush()

        print(json.dumps({"status": "completed"}))

    def _process_txt(self, input_file: Path, output_file: Path):
        with open(input_file, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        anonymized_content = self.anonymizer.anonymize(content, mode=self.mode)
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(anonymized_content)

    _REPLACEMENT_MAP = {
        "PERSON": "[NAME]",
        "PHONE_NUMBER": "[PHONE]",
        "LOCATION": "[ADDRESS]",
        "FISCAL_ID": "[FISCAL_ID]",
        "EMAIL_ADDRESS": "[EMAIL]",
    }

    def _apply_entities_to_runs(self, paragraph, entities) -> None:
        """Apply entity redactions run-by-run to preserve inline formatting."""
        if not paragraph.runs:
            return

        # Build (start, end, run) for each run using original paragraph text
        run_spans: list[tuple[int, int, object]] = []
        pos = 0
        for run in paragraph.runs:
            run_spans.append((pos, pos + len(run.text), run))
            pos += len(run.text)

        # Process right-to-left so earlier positions stay valid
        for entity in sorted(entities, key=lambda e: e.start, reverse=True):
            replacement = self._REPLACEMENT_MAP.get(entity.entity_type, "[REDACTED]")
            ent_start, ent_end = entity.start, entity.end
            placed = False
            for r_start, r_end, run in run_spans:
                if r_end <= ent_start or r_start >= ent_end:
                    continue
                local_start = max(ent_start, r_start) - r_start
                local_end = min(ent_end, r_end) - r_start
                if not placed:
                    run.text = run.text[:local_start] + replacement + run.text[local_end:]
                    placed = True
                else:
                    run.text = run.text[:local_start] + run.text[local_end:]

    def _process_docx(self, input_file: Path, output_file: Path):
        doc = Document(input_file)

        def anonymize_paragraphs(paragraphs):
            for paragraph in paragraphs:
                if not paragraph.text.strip():
                    continue

                entities = self.anonymizer.analyze_and_filter(paragraph.text, mode=self.mode)
                if not entities:
                    continue

                self._apply_entities_to_runs(paragraph, entities)

        anonymize_paragraphs(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    anonymize_paragraphs(cell.paragraphs)
                    
        doc.save(output_file)

    def _process_pdf(self, input_file: Path, output_file: Path):
        doc = fitz.open(input_file)
        try:
            for page in doc:
                text = page.get_text("text")

                filtered = self.anonymizer.analyze_and_filter(text, mode=self.mode)

                # Fallback to OCR if page seems scanned
                if not filtered and len(text.strip()) < 10:
                    render_matrix = fitz.Matrix(2, 2)
                    inv_matrix = ~render_matrix
                    pix = page.get_pixmap(matrix=render_matrix)
                    img = Image.open(io.BytesIO(pix.tobytes()))
                    ocr_data = pytesseract.image_to_data(img, lang='ita', output_type=pytesseract.Output.DICT)
                    ocr_text = " ".join([w for w in ocr_data['text'] if w.strip()])

                    ocr_entities = self.anonymizer.analyze_and_filter(ocr_text, mode=self.mode)

                    for res in ocr_entities:
                        target_words = ocr_text[res.start:res.end].split()
                        for target_word in target_words:
                            for i, word in enumerate(ocr_data['text']):
                                if target_word in word:
                                    x, y, w, h = ocr_data['left'][i], ocr_data['top'][i], ocr_data['width'][i], ocr_data['height'][i]
                                    pixel_rect = fitz.Rect(x, y, x + w, y + h)
                                    rect = pixel_rect * inv_matrix
                                    page.add_redact_annot(rect, fill=(0, 0, 0))
                else:
                    for res in filtered:
                        target_text = text[res.start:res.end]
                        for inst in page.search_for(target_text):
                            page.add_redact_annot(inst, fill=(0, 0, 0))

                page.apply_redactions()

            doc.save(output_file, garbage=4, deflate=True)
        finally:
            doc.close()

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Anonymize documents.")
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--mode", default="default", choices=["default", "aggressive"])
    args = parser.parse_args()
    
    try:
        engine = DocumentAnonymizer()
        processor = FolderProcessor(engine, mode=args.mode)
        processor.process_folder(args.input, args.output)
    except Exception as e:
        print(json.dumps({"status": "error", "message": str(e)}))
        sys.stdout.flush()
        sys.exit(1)
